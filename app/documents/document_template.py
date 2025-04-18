from docx import Document
from docx.shared import Inches
from io import BytesIO
import base64
from datetime import datetime
from sqlalchemy import text
import os

class ProjectDocumentTemplate:
    def __init__(self, project, user, db):
        self.doc = Document()
        self.project = project
        self.user = user
        self.db = db

    def add_cover_page(self, logo_base64: str = None):
        if logo_base64:
            self._insert_logo(logo_base64)
        else:
            self.doc.add_paragraph("FounderHub", style="Title")
            self.doc.add_paragraph("Upgrade to Pro for AI branding & logo support.", style="Intense Quote")

        self.doc.add_paragraph(self.project.title, style="Title")
        self.doc.add_paragraph(f"Prepared for: {self.user['name']} ({self.user['email']})")
        self.doc.add_paragraph(f"Project Type: {self.project.type or 'General'}")
        self.doc.add_paragraph("Date: " + datetime.utcnow().strftime("%B %d, %Y"))
        self.doc.add_paragraph("\n")

    def add_table_of_contents(self):
        self.doc.add_paragraph("Table of Contents", style="Heading 1")
        toc = [
            "1. Executive Summary",
            "2. Problem & Solution",
            "3. Market & Opportunity",
            "4. Product Vision",
            "5. Team & Advisors",
            "6. Go-to-Market Strategy",
            "7. Financial Overview",
            "8. Risks & Mitigations",
            "9. Board Decision",
            "10. Appendices"
        ]
        for item in toc:
            self.doc.add_paragraph(item)

    def add_section(self, heading: str, content: str):
        self.doc.add_heading(heading, level=1)
        self.doc.add_paragraph(content)

    def _insert_logo(self, base64_str):
        image_stream = BytesIO(base64.b64decode(base64_str))
        self.doc.add_picture(image_stream, width=Inches(2))

    def save(self, output_path: str):
        self.doc.save(output_path)


def get_project_logo_base64(project_id: str, db):
    row = db.execute(
        text("""
        SELECT file_base64 FROM project_assets
        WHERE project_id = :project_id AND type = 'logo'
        ORDER BY updated_at DESC LIMIT 1
        """),
        {"project_id": project_id}
    ).fetchone()

    if row:
        return row.file_base64

    fallback_path = "app/static/logo.png"
    if os.path.exists(fallback_path):
        with open(fallback_path, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")
    return None


def build_project_document(project, user, db, sections: dict, output_path: str):
    logo_base64 = get_project_logo_base64(str(project.id), db)
    template = ProjectDocumentTemplate(project, user, db)
    template.add_cover_page(logo_base64)
    template.add_table_of_contents()

    for heading, content in sections.items():
        template.add_section(heading, content)

    template.save(output_path)
    return output_path