import os
import tempfile
import base64
import re
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML, CSS
from msal import ConfidentialClientApplication
from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError
import requests

# === Directories ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "../../templates")
STATIC_DIR = os.path.join(BASE_DIR, "../../static")

# === Graph Config ===
MS_TOKEN_URL = f"https://login.microsoftonline.com/{os.getenv('MS_TENANT_ID')}/oauth2/v2.0/token"
GRAPH_API_SENDMAIL = f"https://graph.microsoft.com/v1.0/users/{os.getenv('MS_SENDER_EMAIL')}/sendMail"

# === Clean Markdown Formatting ===
def clean_text(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # remove bold
    text = re.sub(r"(?<!\n)\s*(\d+\.)", r"\n\1", text)  # add line breaks before numbered items
    text = re.sub(r"(?<!\n)\s*(- )", r"\n\1", text)  # line breaks before dashes
    text = re.sub(r"&nbsp;", " ", text)
    return text.strip()

# === PDF Generator ===
def render_pdf(data: dict) -> str:
    # Clean all messages
    for msg in data["messages"]:
        msg["message"] = clean_text(msg["message"])

    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))
    template = env.get_template("pdf_template.html")
    html_content = template.render(**data)

    css = CSS(string=f"""
        @page {{
            size: A4;
            margin: 1in;
            @top-center {{
                content: "FounderHub.ai – {data['idea'].title}";
                font-size: 10pt;
                color: #888;
            }}
            @bottom-center {{
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
                color: #888;
            }}
        }}

        body {{
            font-family: 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.6;
        }}

        h1, h2, h3 {{
            color: #0F172A;
        }}

        .chat-role {{
            font-weight: bold;
            margin-top: 1em;
            color: #0F172A;
        }}

        .chat-message {{
            margin-bottom: 1em;
            white-space: pre-wrap;
        }}
    """)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    HTML(string=html_content, base_url=STATIC_DIR).write_pdf(tmp.name, stylesheets=[css])
    return tmp.name

# === DOCX Generator ===
def render_docx(data: dict) -> str:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.75)

    # Logo
    logo_path = os.path.join(STATIC_DIR, "logo.png")
    try:
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(1.5))
    except UnrecognizedImageError:
        print("⚠️ Invalid logo format.")

    doc.add_heading("FounderHub Deep Dive Report", level=1)
    doc.add_paragraph(f"Date: {data['date']}")

    doc.add_heading("🧠 Idea Summary", level=2)
    doc.add_paragraph(f"Title: {data['idea'].title}")
    doc.add_paragraph(f"Problem: {data['idea'].problem}")
    doc.add_paragraph(f"Audience: {data['idea'].audience}")
    doc.add_paragraph(f"Solution: {data['idea'].solution}")
    doc.add_paragraph(f"Score: {data['score']}/100")

    doc.add_page_break()
    doc.add_heading("💬 Chat Log", level=2)
    for msg in data["messages"]:
        doc.add_paragraph(f"{msg['role'].capitalize()}:", style="Heading 3")
        doc.add_paragraph(clean_text(msg["message"]))

    doc.add_page_break()
    doc.add_heading("📊 Stats", level=2)
    doc.add_paragraph(f"Tokens Used: {data['tokens_used']} / {data['token_limit']}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name

# === Microsoft Graph Auth ===
def get_graph_token():
    client_id = os.getenv("MS_CLIENT_ID")
    client_secret = os.getenv("MS_CLIENT_SECRET")
    tenant_id = os.getenv("MS_TENANT_ID")
    authority = f"https://login.microsoftonline.com/{tenant_id}"

    app = ConfidentialClientApplication(
        client_id=client_id,
        client_credential=client_secret,
        authority=authority
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" not in result:
        raise Exception(f"Graph token error: {result.get('error_description', 'Unknown error')}")
    return result["access_token"]

# === Send Email via Graph ===
def send_email_with_attachment(to_email: str, subject: str, body: str, file_path: str):
    access_token = get_graph_token()
    with open(file_path, "rb") as f:
        content_bytes = base64.b64encode(f.read()).decode("utf-8")

    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_name)[1][1:].lower()
    content_type = (
        "application/pdf"
        if ext == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    mail = {
        "message": {
            "subject": subject,
            "body": {"contentType": "HTML", "content": body},
            "toRecipients": [{"emailAddress": {"address": to_email}}],
            "attachments": [{
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": file_name,
                "contentType": content_type,
                "contentBytes": content_bytes
            }]
        },
        "saveToSentItems": "true"
    }

    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }

    response = requests.post(GRAPH_API_SENDMAIL, headers=headers, json=mail)
    if response.status_code >= 400:
        raise Exception(f"Graph API Error {response.status_code}: {response.text}")
